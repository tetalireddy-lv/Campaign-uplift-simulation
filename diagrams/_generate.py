"""Generate Architecture.png and Interaction_of_Agents.docx for the diagrams folder."""
import os
import sys
import textwrap

# ── Architecture.png ──────────────────────────────────────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── colour palette ────────────────────────────────────────────────────────────
C = {
    "user":     "#4F46E5",   # indigo  – browser / user
    "frontend": "#0EA5E9",   # sky     – React/Vite
    "backend":  "#10B981",   # emerald – FastAPI
    "agent":    "#F59E0B",   # amber   – LangGraph nodes
    "tool":     "#EF4444",   # red     – tool functions
    "llm":      "#8B5CF6",   # violet  – Azure OpenAI
    "store":    "#6B7280",   # gray    – session store / mock
    "arrow":    "#374151",   # dark
    "bg":       "#F9FAFB",
    "white":    "#FFFFFF",
}

fig, ax = plt.subplots(figsize=(22, 14))
ax.set_xlim(0, 22)
ax.set_ylim(0, 14)
ax.set_facecolor(C["bg"])
fig.patch.set_facecolor(C["bg"])
ax.axis("off")


# ─── helpers ──────────────────────────────────────────────────────────────────
def box(ax, x, y, w, h, label, sublabel="", color="#FFFFFF", border="#000", fontsize=9, radius=0.3):
    patch = FancyBboxPatch((x, y), w, h,
                           boxstyle=f"round,pad=0,rounding_size={radius}",
                           linewidth=1.5, edgecolor=border, facecolor=color, zorder=3)
    ax.add_patch(patch)
    if sublabel:
        ax.text(x + w/2, y + h/2 + 0.15, label, ha="center", va="center",
                fontsize=fontsize, fontweight="bold", color=C["arrow"], zorder=4)
        ax.text(x + w/2, y + h/2 - 0.22, sublabel, ha="center", va="center",
                fontsize=fontsize - 1, color="#6B7280", zorder=4)
    else:
        ax.text(x + w/2, y + h/2, label, ha="center", va="center",
                fontsize=fontsize, fontweight="bold", color=C["arrow"], zorder=4)


def section_bg(ax, x, y, w, h, label, color, alpha=0.08):
    patch = FancyBboxPatch((x, y), w, h,
                           boxstyle="round,pad=0,rounding_size=0.4",
                           linewidth=2, edgecolor=color, facecolor=color,
                           alpha=alpha, zorder=1)
    ax.add_patch(patch)
    ax.text(x + 0.2, y + h - 0.25, label, fontsize=8, color=color,
            fontweight="bold", va="top", zorder=2)


def arrow(ax, x1, y1, x2, y2, label="", color=C["arrow"]):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=1.4,
                                connectionstyle="arc3,rad=0.0"), zorder=5)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx + 0.1, my, label, fontsize=7, color=color, va="center")


def curved_arrow(ax, x1, y1, x2, y2, rad=0.3, label="", color=C["arrow"]):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=1.4,
                                connectionstyle=f"arc3,rad={rad}"), zorder=5)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx + 0.15, my + 0.1, label, fontsize=7, color=color)


# ─── Title ────────────────────────────────────────────────────────────────────
ax.text(11, 13.5, "Campaign Readiness Copilot — System Architecture",
        ha="center", va="top", fontsize=15, fontweight="bold", color=C["arrow"])

# ════════════════════════════════════════════════════════════════════════
# ROW 1 – User + Frontend (top-left band)
# ════════════════════════════════════════════════════════════════════════
section_bg(ax, 0.3, 10.8, 6.8, 2.3, "Presentation Layer", C["frontend"])
box(ax, 0.6, 11.5, 1.8, 1.2, "Campaign\nManager", color="#EFF6FF", border=C["user"], fontsize=8)
arrow(ax, 2.4, 12.1, 3.2, 12.1, "HTTP")
box(ax, 3.2, 11.0, 3.5, 2.2, "React + Vite", "localhost:5173",
    color="#F0F9FF", border=C["frontend"])

# ════════════════════════════════════════════════════════════════════════
# ROW 1 – Backend (top-right band)
# ════════════════════════════════════════════════════════════════════════
section_bg(ax, 7.4, 10.8, 6.5, 2.3, "API Layer  (FastAPI · localhost:8000)", C["backend"])

api_routes = [
    ("POST /start", 7.7, 11.5),
    ("POST /validate-readiness", 9.2, 11.5),
    ("POST /resolve-ambiguity", 10.8, 11.5),
    ("POST /approve-assumptions", 12.2, 11.5),
]
for label, bx, by in api_routes:
    box(ax, bx, by, 1.35, 0.7, label, color="#ECFDF5", border=C["backend"], fontsize=7)

box(ax, 9.5, 10.95, 1.35, 0.7, "POST /plan-simulate", color="#ECFDF5", border=C["backend"], fontsize=7)
box(ax, 11.0, 10.95, 1.35, 0.7, "POST /qa-handoff", color="#ECFDF5", border=C["backend"], fontsize=7)

# Frontend → Backend
arrow(ax, 6.7, 12.1, 7.4, 12.1, "REST /api/*")

# Session Store (top-right corner)
section_bg(ax, 14.1, 10.8, 2.8, 2.3, "Session Store", C["store"])
box(ax, 14.3, 11.4, 2.4, 1.3, "In-Memory\nSession Store", "session.py",
    color="#F9FAFB", border=C["store"])
arrow(ax, 13.9, 12.0, 14.3, 12.0)

# ════════════════════════════════════════════════════════════════════════
# ROW 2 – LangGraph Workflow (centre)
# ════════════════════════════════════════════════════════════════════════
section_bg(ax, 0.3, 5.5, 16.6, 5.0, "Agent Workflow Layer  (LangGraph)", C["agent"])

nodes = [
    ("parse_brief\n(Step 1)", 0.7, 8.8),
    ("validate_readiness\n(Step 2)", 3.0, 8.8),
    ("resolve_ambiguity\n(Step 3)", 5.5, 8.8),
    ("await_manager\napproval (3b)", 8.0, 8.8),
    ("generate_plan\n(Step 4)", 10.5, 8.8),
    ("qa_and_handoff\n(Step 5)", 13.0, 8.8),
    ("final_approval", 15.5, 8.8),
]
node_w, node_h = 2.0, 0.9
for label, nx, ny in nodes:
    box(ax, nx, ny, node_w, node_h, label, color="#FFFBEB", border=C["agent"], fontsize=7.5)

# Arrows between nodes
pairs = [
    (0.7+node_w, 8.8+node_h/2, 3.0, 8.8+node_h/2, ""),
    (3.0+node_w, 8.8+node_h/2, 5.5, 8.8+node_h/2, "score<85"),
    (5.5+node_w, 8.8+node_h/2, 8.0, 8.8+node_h/2, ""),
    (8.0+node_w, 8.8+node_h/2, 10.5, 8.8+node_h/2, ""),
    (10.5+node_w, 8.8+node_h/2, 13.0, 8.8+node_h/2, ""),
    (13.0+node_w, 8.8+node_h/2, 15.5, 8.8+node_h/2, "QA pass"),
]
for x1,y1,x2,y2,lbl in pairs:
    arrow(ax, x1, y1, x2, y2, lbl, C["agent"])

# Skip-ambiguity dashed bypass (validate → generate_plan if score>=85)
ax.annotate("", xy=(10.5+0.2, 8.8+node_h), xytext=(3.0+node_w/2, 8.8+node_h),
            arrowprops=dict(arrowstyle="-|>", color=C["agent"], lw=1.2, linestyle="dashed",
                            connectionstyle="arc3,rad=-0.3"), zorder=5)
ax.text(8.0, 10.0, "score≥85 (skip ambiguity)", fontsize=7, color=C["agent"], ha="center", style="italic")

# QA loop-back (qa_and_handoff → generate_plan if critical misalignments)
ax.annotate("", xy=(10.5+1.0, 8.8), xytext=(13.0+1.0, 8.8),
            arrowprops=dict(arrowstyle="-|>", color="#DC2626", lw=1.2, linestyle="dashed",
                            connectionstyle="arc3,rad=0.4"), zorder=5)
ax.text(11.8, 7.95, "QA fail → re-plan", fontsize=7, color="#DC2626", ha="center", style="italic")

# Backend → LangGraph entry
arrow(ax, 7.0, 11.0, 1.7, 9.7, "", C["backend"])
ax.text(4.3, 10.5, "invoke workflow", fontsize=7, color=C["backend"], style="italic")

# ── Tool boxes (below nodes) ──────────────────────────────────────────────────
section_bg(ax, 0.3, 5.6, 16.6, 2.9, "Tool Functions (14 modules)", C["tool"])

tool_rows = [
    # row 1
    [("brief_parser", 0.6, 7.6), ("classify_intent", 2.4, 7.6), ("gap_detector", 4.2, 7.6),
     ("compliance_risk", 6.0, 7.6), ("kpi_review", 7.8, 7.6), ("clarification_q", 9.6, 7.6),
     ("assumption_reg", 11.4, 7.6)],
    # row 2
    [("channel_strategy", 0.6, 6.7), ("execution_plan", 2.7, 6.7), ("asset_checklist", 4.8, 6.7),
     ("timeline_wb", 6.9, 6.7), ("uplift_sim", 9.0, 6.7), ("brief_to_plan_qa", 11.1, 6.7),
     ("consistency_chk", 13.2, 6.7), ("handoff_packet", 15.3, 6.7)],
]
tw, th = 1.75, 0.65
for row in tool_rows:
    for label, tx, ty in row:
        box(ax, tx, ty, tw, th, label, color="#FFF1F2", border=C["tool"], fontsize=7)

# Node → tool arrows (representative, not all)
node_tool_links = [
    (0.7+node_w/2, 8.8, 0.6+tw/2, 7.6+th),     # parse_brief → brief_parser
    (0.7+node_w/2, 8.8, 2.4+tw/2, 7.6+th),      # parse_brief → classify_intent
    (3.0+node_w/2, 8.8, 4.2+tw/2, 7.6+th),      # validate → gap_detector
    (3.0+node_w/2, 8.8, 6.0+tw/2, 7.6+th),      # validate → compliance_risk
    (3.0+node_w/2, 8.8, 7.8+tw/2, 7.6+th),      # validate → kpi_review
    (5.5+node_w/2, 8.8, 9.6+tw/2, 7.6+th),      # resolve → clarification_q
    (5.5+node_w/2, 8.8, 11.4+tw/2, 7.6+th),     # resolve → assumption_reg
    (10.5+node_w/2, 8.8, 0.6+tw/2, 6.7+th),     # generate_plan → channel_strategy
    (10.5+node_w/2, 8.8, 9.0+tw/2, 6.7+th),     # generate_plan → uplift_sim
    (13.0+node_w/2, 8.8, 11.1+tw/2, 6.7+th),    # qa → brief_to_plan_qa
    (13.0+node_w/2, 8.8, 15.3+tw/2, 6.7+th),    # qa → handoff_packet
]
for x1, y1, x2, y2 in node_tool_links:
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=C["tool"], lw=0.8,
                                connectionstyle="arc3,rad=0.0"), zorder=4)

# ════════════════════════════════════════════════════════════════════════
# ROW 3 – Azure OpenAI + Mock Data (bottom)
# ════════════════════════════════════════════════════════════════════════
section_bg(ax, 0.3, 0.3, 10.0, 4.9, "LLM Provider", C["llm"])
box(ax, 0.7, 2.2, 4.0, 2.5, "Azure OpenAI\n(live mode)", "GPT deployment via\nllm_utils.py",
    color="#F5F3FF", border=C["llm"])
box(ax, 5.2, 2.2, 4.5, 2.5, "Jinja2 Prompt Templates\n(.j2 files)", "14 templates\norganised by step",
    color="#F5F3FF", border=C["llm"])

arrow(ax, 4.7, 3.4, 5.2, 3.4, "render")
# tools → llm
arrow(ax, 8.0, 6.7, 6.5, 4.7, "", C["llm"])
ax.text(7.6, 5.7, "LLM calls", fontsize=7, color=C["llm"], style="italic")

section_bg(ax, 10.6, 0.3, 6.4, 4.9, "Fallback / Data", C["store"])
box(ax, 10.9, 2.8, 2.8, 1.8, "Mock Data\nFallback", "mock_data.py",
    color="#F3F4F6", border=C["store"])
box(ax, 14.0, 2.8, 2.7, 1.8, "CSV Datasets", "campaign_briefs\nhistorical_results",
    color="#F3F4F6", border=C["store"])
box(ax, 10.9, 0.5, 5.8, 1.9, "Domain Knowledge Markdown",
    "campaign_domain · brief_schema · compliance · kpi ·\nchannel_strategy · uplift_simulation · approval_workflow",
    color="#F3F4F6", border=C["store"])

# tools → mock
arrow(ax, 8.0, 6.7, 11.8, 4.6, "", C["store"])

# ─── Legend ───────────────────────────────────────────────────────────────────
legend_items = [
    (C["frontend"], "Frontend (React/Vite)"),
    (C["backend"],  "Backend (FastAPI)"),
    (C["agent"],    "LangGraph Nodes"),
    (C["tool"],     "Tool Functions"),
    (C["llm"],      "LLM / Prompts"),
    (C["store"],    "Session / Data"),
]
lx, ly = 17.2, 13.2
ax.text(lx, ly, "Legend", fontsize=9, fontweight="bold", color=C["arrow"])
for i, (color, label) in enumerate(legend_items):
    ry = ly - 0.55 * (i+1)
    patch = FancyBboxPatch((lx, ry-0.17), 0.4, 0.34,
                           boxstyle="round,pad=0,rounding_size=0.06",
                           edgecolor=color, facecolor=color, alpha=0.3, zorder=6)
    ax.add_patch(patch)
    ax.text(lx + 0.55, ry, label, fontsize=8, color=C["arrow"], va="center", zorder=7)

plt.tight_layout(pad=0)
arch_path = os.path.join(OUT_DIR, "Architecture.png")
plt.savefig(arch_path, dpi=180, bbox_inches="tight", facecolor=C["bg"])
plt.close()
print(f"Saved: {arch_path}")


# ── Interaction_of_Agents.docx ────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
section = doc.sections[0]
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)

BRAND   = RGBColor(0x4F, 0x46, 0xE5)   # indigo
AMBER   = RGBColor(0xF5, 0x9E, 0x0B)
RED     = RGBColor(0xEF, 0x44, 0x44)
GREEN   = RGBColor(0x10, 0xB9, 0x81)
GRAY    = RGBColor(0x6B, 0x72, 0x80)
BLACK   = RGBColor(0x11, 0x18, 0x27)

def heading(doc, text, level=1, color=BRAND):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p

def para(doc, text, bold=False, italic=False, color=None, size=11, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "4F46E5")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx+1].cells
        fill = "F0F9FF" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            for run in row[c_idx].paragraphs[0].runs:
                run.font.size = Pt(9.5)
            tc = row[c_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def divider(doc):
    p = doc.add_paragraph("─" * 90)
    p.runs[0].font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
    p.runs[0].font.size = Pt(8)


# ════════════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Campaign Readiness Copilot")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = BRAND

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Agent Interaction & Workflow Reference")
run2.font.size = Pt(14)
run2.font.color.rgb = GRAY
run2.italic = True

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(
    "This document describes every agent node, tool function, and data-flow\n"
    "in the LangGraph-powered campaign readiness workflow."
)
run3.font.size = Pt(11)
run3.font.color.rgb = BLACK

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ════════════════════════════════════════════════════════════════════════
heading(doc, "1. System Overview")
para(doc,
     "The Campaign Readiness Copilot is a full-stack AI application that guides a campaign "
     "manager from a raw campaign brief to a launch-ready execution plan. It combines a "
     "FastAPI backend, a React/Vite frontend, and a LangGraph-orchestrated multi-step agent "
     "workflow powered by Azure OpenAI. The AI never launches campaigns autonomously — every "
     "consequential decision requires explicit manager approval.")

doc.add_paragraph()
heading(doc, "Key Design Principles", level=2, color=AMBER)
bullets = [
    "Human-in-the-loop: manager approves assumptions before planning proceeds.",
    "Graceful degradation: every tool falls back to realistic mock data when no Azure "
    "OpenAI credentials are present.",
    "Parallel tool execution: independent tools within each step run concurrently (asyncio) "
    "to minimise latency.",
    "Conditional routing: the LangGraph graph skips or replays nodes based on readiness "
    "score and QA outcomes.",
    "Session isolation: each workflow run lives in its own in-memory session keyed by UUID.",
]
for b in bullets:
    bullet(doc, b)

divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 2. AGENT NODES
# ════════════════════════════════════════════════════════════════════════
heading(doc, "2. Agent Nodes (LangGraph)")
para(doc,
     "The workflow is modelled as a directed graph with seven nodes. Each node is a Python "
     "async function that reads from and writes to the shared CampaignReadinessState object.")

doc.add_paragraph()
nodes_data = [
    ("parse_brief",            "Step 1", "Parses raw brief text → structured brief + campaign intent",
     "brief_parser, classify_intent", "Always"),
    ("validate_readiness",     "Step 2", "Scores completeness; identifies gaps, compliance risks, KPI gaps",
     "gap_detector, compliance_risk, kpi_review", "After parse_brief"),
    ("resolve_ambiguity",      "Step 3", "Generates clarifying questions + assumption register",
     "clarification_questions, assumption_register", "If readiness score < 85"),
    ("await_manager_approval", "Step 3b","Pauses graph; manager reviews and approves/edits assumptions",
     "None (human gate)", "After resolve_ambiguity"),
    ("generate_plan",          "Step 4", "Builds orchestrator plan → channel strategy → execution plan "
     "→ asset checklist → timeline → uplift simulation",
     "channel_strategy, execution_plan, asset_checklist,\ntimeline_workback, uplift_sim", "Always"),
    ("qa_and_handoff",         "Step 5", "Runs three QA checks concurrently; generates handoff packet",
     "brief_to_plan_qa, multi_channel_consistency,\nfinal_compliance_review, handoff_packet",
     "After generate_plan"),
    ("final_approval",         "End",    "Records final status; session transitions to READY_FOR_REVIEW",
     "None", "After successful QA"),
]
add_table(doc,
    ["Node", "Step", "Responsibility", "Tools Invoked", "Triggered When"],
    nodes_data,
    col_widths=[1.3, 0.55, 2.3, 2.15, 1.4],
)
doc.add_paragraph()
divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 3. CONDITIONAL ROUTING
# ════════════════════════════════════════════════════════════════════════
heading(doc, "3. Conditional Routing Logic")
para(doc, "Two conditional edge functions control non-linear flow in the graph:")

doc.add_paragraph()
heading(doc, "3.1  should_skip_ambiguity  (after validate_readiness)", level=2, color=GREEN)
para(doc, "Reads: state['readiness_score']")
bullet(doc, "score ≥ 85  →  jump directly to generate_plan  (brief is high quality, no clarification needed)")
bullet(doc, "score < 85  →  proceed to resolve_ambiguity  (gaps or risks detected)")
doc.add_paragraph()

heading(doc, "3.2  should_proceed_after_qa  (after qa_and_handoff)", level=2, color=RED)
para(doc, "Reads: state['qa_report']['critical_misalignments']")
bullet(doc, "[] (empty list)  →  proceed to final_approval  (plan is consistent)")
bullet(doc, "Non-empty  →  loop back to generate_plan  (critical issues found; re-plan with QA context)")
doc.add_paragraph()
divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 4. TOOL FUNCTIONS
# ════════════════════════════════════════════════════════════════════════
heading(doc, "4. Tool Functions")
para(doc,
     "Each tool is an async Python function in agent/tools/. "
     "It renders a Jinja2 prompt template, calls the LLM (or falls back to mock data), "
     "and returns a typed dict consumed by its parent node.")

doc.add_paragraph()
tools_data = [
    # Step 1
    ("brief_parser",               "Step 1", "Parse Brief",
     "Extracts objectives, target audience, budget, dates, channels, KPIs, and constraints "
     "from free-form text.", "brief_parser.j2"),
    ("classify_intent",            "Step 1", "Parse Brief",
     "Classifies campaign type (awareness, conversion, retention …) and primary goal.",
     "classify_campaign_intent.j2"),
    # Step 2
    ("gap_detector",               "Step 2", "Validate Readiness",
     "Identifies missing mandatory fields; scores brief completeness (0–100).",
     "gap_detector.j2"),
    ("compliance_risk",            "Step 2", "Validate Readiness",
     "Flags regulatory, brand, and legal risks based on channel + audience combination.",
     "compliance_risk_review.j2"),
    ("kpi_review",                 "Step 2", "Validate Readiness",
     "Assesses whether stated KPIs are measurable and aligned with campaign type.",
     "kpi_measurement_review.j2"),
    # Step 3
    ("clarification_questions",    "Step 3", "Resolve Ambiguity",
     "Generates prioritised clarifying questions for the campaign manager.",
     "clarification_questions.j2"),
    ("assumption_register",        "Step 3", "Resolve Ambiguity",
     "Creates a register of AI-generated assumptions to fill brief gaps; "
     "awaits manager approval before planning.",
     "assumption_register.j2"),
    # Step 4
    ("channel_strategy",           "Step 4", "Plan & Simulate",
     "Selects optimal channel mix, allocates budget percentages, sets channel-level KPI targets.",
     "channel_strategy.j2"),
    ("execution_plan_generator",   "Step 4", "Plan & Simulate",
     "Builds a week-by-week execution plan with owners, deliverables, and dependencies.",
     "execution_plan_generator.j2"),
    ("asset_checklist",            "Step 4", "Plan & Simulate",
     "Produces a complete creative asset checklist mapped to channels and launch dates.",
     "asset_checklist.j2"),
    ("timeline_workback",          "Step 4", "Plan & Simulate",
     "Generates a reverse-engineered timeline from launch date; flags critical-path items.",
     "timeline_workback.j2"),
    ("campaign_uplift_simulator",  "Step 4", "Plan & Simulate",
     "Runs a statistical uplift simulation using historical segment data; "
     "returns expected lift, confidence interval, and sensitivity analysis.",
     "market_uplift_simulator.j2"),
    # Step 5
    ("brief_to_plan_qa",           "Step 5", "QA & Handoff",
     "Cross-checks the execution plan against the original brief for objective alignment.",
     "brief_to_plan_qa.j2"),
    ("multi_channel_consistency",  "Step 5", "QA & Handoff",
     "Verifies message consistency, audience overlap, and budget coherence across channels.",
     "multi_channel_consistency.j2"),
    ("final_compliance_review",    "Step 5", "QA & Handoff",
     "Final regulatory and brand safety pass over the complete plan.",
     "final_compliance_review.j2"),
    ("stakeholder_handoff_packet", "Step 5", "QA & Handoff",
     "Assembles executive summary, risk register, next-steps, and owner matrix for stakeholders.",
     "stakeholder_handoff_packet.j2"),
]
add_table(doc,
    ["Tool", "Step", "Phase", "What It Does", "Prompt Template"],
    tools_data,
    col_widths=[1.5, 0.5, 1.0, 3.1, 1.65],
)
doc.add_paragraph()
divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 5. API → WORKFLOW MAPPING
# ════════════════════════════════════════════════════════════════════════
heading(doc, "5. REST API → Workflow Mapping")
para(doc,
     "The FastAPI backend exposes seven routes. Each route triggers one workflow function "
     "in workflow_service.py which, in turn, calls the relevant tool(s).")

doc.add_paragraph()
api_data = [
    ("GET  /api/config/status",                 "—",                        "Reports live vs mock LLM mode"),
    ("POST /api/workflow/start",                "parse_brief node",         "Runs brief_parser + classify_intent concurrently"),
    ("POST /api/workflow/{id}/validate-readiness","validate_readiness node", "Runs gap_detector, compliance_risk, kpi_review concurrently"),
    ("POST /api/workflow/{id}/resolve-ambiguity","resolve_ambiguity node",  "Sequential: clarification_questions then assumption_register"),
    ("POST /api/workflow/{id}/approve-assumptions","await_manager_approval","Manager submits approved assumptions; no LLM call"),
    ("POST /api/workflow/{id}/plan-and-simulate","generate_plan node",      "Orchestrator → channel → plan (seq); asset/timeline/sim (parallel)"),
    ("POST /api/workflow/{id}/qa-and-handoff",  "qa_and_handoff node",      "brief_to_plan_qa, consistency, compliance (parallel) → handoff (seq)"),
    ("GET  /api/workflow/{id}",                 "—",                        "Returns full session state snapshot"),
]
add_table(doc,
    ["Endpoint", "Workflow Node", "Execution Pattern"],
    api_data,
    col_widths=[2.4, 1.8, 3.55],
)
doc.add_paragraph()
divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 6. STATE OBJECT
# ════════════════════════════════════════════════════════════════════════
heading(doc, "6. Shared State  (CampaignReadinessState)")
para(doc,
     "All nodes communicate exclusively through a typed TypedDict. "
     "No direct function-to-function calls — the graph passes state between nodes.")

doc.add_paragraph()
state_fields = [
    ("raw_brief",             "str",         "Original unstructured brief text"),
    ("structured_brief",      "dict",         "Parsed brief fields (objective, audience, budget …)"),
    ("campaign_intent",       "dict",         "Classified campaign type and primary goal"),
    ("readiness_score",       "int (0–100)",  "Completeness score; drives routing decision"),
    ("gap_report",            "dict",         "Missing fields, risk flags, recommended fixes"),
    ("compliance_report",     "dict",         "Regulatory and brand risk flags with severity"),
    ("kpi_report",            "dict",         "KPI measurability assessment"),
    ("clarification_questions","list[dict]",  "Prioritised questions for the campaign manager"),
    ("assumption_register",   "list[dict]",   "AI-generated assumptions awaiting approval"),
    ("approved_assumptions",  "list[dict]",   "Manager-approved assumptions used in planning"),
    ("manager_answers",       "dict",         "Optional free-text answers from the manager"),
    ("orchestrator_plan",     "dict",         "High-level strategic direction for the plan"),
    ("channel_strategy",      "dict",         "Channel mix, budget split, channel KPIs"),
    ("execution_plan",        "dict",         "Week-by-week plan with owners and deliverables"),
    ("asset_checklist",       "list[dict]",   "Creative assets needed per channel"),
    ("timeline",              "list[dict]",   "Reverse-engineered launch timeline"),
    ("measurement_plan",      "dict",         "Tracking setup and measurement methodology"),
    ("simulation_report",     "dict",         "Uplift estimate, confidence interval, sensitivity"),
    ("qa_report",             "dict",         "QA findings; critical_misalignments drives re-plan loop"),
    ("compliance_final",      "dict",         "Final compliance pass outcome"),
    ("handoff_packet",        "dict",         "Complete stakeholder-ready handoff document"),
    ("status",                "str (enum)",   "Current workflow status (e.g. PLAN_GENERATED)"),
    ("audit_log",             "list[dict]",   "Timestamped record of every node execution"),
]
add_table(doc,
    ["Field", "Type", "Description"],
    state_fields,
    col_widths=[1.85, 1.35, 4.55],
)
doc.add_paragraph()
divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 7. DATA FLOW NARRATIVE
# ════════════════════════════════════════════════════════════════════════
heading(doc, "7. End-to-End Data Flow")

steps = [
    ("1 → Brief Submitted",
     "The campaign manager pastes a raw campaign brief into the React UI. "
     "The frontend sends POST /api/workflow/start to the FastAPI backend."),
    ("2 → Parse Brief",
     "workflow_service.run_parse_brief() runs brief_parser and classify_intent concurrently. "
     "Results are stored in the session as structured_brief and campaign_intent. "
     "Session status → BRIEF_PARSED."),
    ("3 → Validate Readiness",
     "POST /validate-readiness triggers gap_detector, compliance_risk, and kpi_review in parallel. "
     "A composite readiness_score (0–100) is computed and stored. Status → READINESS_VALIDATED."),
    ("4 → Routing Decision",
     "The graph evaluates readiness_score. "
     "≥ 85: jumps directly to generate_plan. "
     "< 85: proceeds to resolve_ambiguity."),
    ("5 → Resolve Ambiguity (conditional)",
     "clarification_questions runs first; its output seeds assumption_register (sequential dependency). "
     "Status → AWAITING_MANAGER_APPROVAL."),
    ("6 → Manager Approval Gate",
     "The UI presents assumptions and questions to the campaign manager. "
     "POST /approve-assumptions stores approved_assumptions and optional manager_answers. "
     "Status → ASSUMPTIONS_APPROVED."),
    ("7 → Plan & Simulate",
     "Sequential: orchestrator_plan → channel_strategy → execution_plan. "
     "Parallel (after execution_plan): asset_checklist, timeline_workback, measurement_plan. "
     "Sequential (after measurement_plan): uplift simulation. Status → PLAN_GENERATED."),
    ("8 → QA & Handoff",
     "brief_to_plan_qa, multi_channel_consistency, and final_compliance_review run concurrently. "
     "If critical_misalignments is non-empty, the graph loops back to generate_plan with QA context. "
     "Otherwise handoff_packet is generated. Status → READY_FOR_REVIEW."),
    ("9 → Handoff",
     "The frontend renders the complete plan, uplift simulation, and handoff packet. "
     "The manager downloads or shares the stakeholder-ready document."),
]

for title, text in steps:
    p = doc.add_paragraph()
    run_t = p.add_run(title + "  ")
    run_t.bold = True
    run_t.font.color.rgb = BRAND
    run_t.font.size = Pt(11)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10.5)

doc.add_paragraph()
divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 8. CONCURRENCY MODEL
# ════════════════════════════════════════════════════════════════════════
heading(doc, "8. Concurrency Model")
para(doc,
     "All tool functions are async. workflow_service.py uses asyncio.gather() to run "
     "independent tools concurrently within a step, reducing total round-trip time. "
     "Sequential dependencies (e.g. assumption_register needs clarification_questions output, "
     "or uplift_sim needs measurement_plan) are enforced by awaiting intermediate results before "
     "launching dependent calls.")

doc.add_paragraph()
conc_data = [
    ("Step 1 – Parse Brief",       "brief_parser + classify_intent",                              "Parallel"),
    ("Step 2 – Validate",          "gap_detector + compliance_risk + kpi_review",                 "Parallel"),
    ("Step 3 – Ambiguity",         "clarification_questions → assumption_register",                "Sequential"),
    ("Step 4 – Plan (phase A)",    "orchestrator_plan → channel_strategy → execution_plan",        "Sequential"),
    ("Step 4 – Plan (phase B)",    "asset_checklist + timeline_workback + measurement_plan",       "Parallel (after exec_plan)"),
    ("Step 4 – Simulate",          "uplift_sim",                                                   "Sequential (after measurement_plan)"),
    ("Step 5 – QA",                "brief_to_plan_qa + multi_channel_consistency + final_compliance","Parallel"),
    ("Step 5 – Handoff",           "handoff_packet",                                               "Sequential (after all QA)"),
]
add_table(doc,
    ["Step", "Tools", "Execution"],
    conc_data,
    col_widths=[1.8, 3.8, 2.15],
)
doc.add_paragraph()
divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 9. FALLBACK STRATEGY
# ════════════════════════════════════════════════════════════════════════
heading(doc, "9. Mock / Fallback Strategy")
para(doc,
     "Each tool checks AZURE_OPENAI_API_KEY at call time. "
     "If the key is absent or matches the placeholder value, the tool returns a pre-built "
     "realistic response from mock_data.py instead of calling the LLM. "
     "This means the entire 5-step workflow runs end-to-end without any API credentials — "
     "useful for demos, CI testing, and local development.")

doc.add_paragraph()
bullet(doc, "GET /api/config/status exposes llm_mode: 'live' | 'mock' so the UI can show a banner.")
bullet(doc, "Mock responses are schema-identical to live responses — the frontend is unaware of the difference.")
bullet(doc, "CSV datasets in agent/data/raw/ provide realistic historical data for the uplift simulator in mock mode.")

doc.add_paragraph()
divider(doc)


# ════════════════════════════════════════════════════════════════════════
# 10. GLOSSARY
# ════════════════════════════════════════════════════════════════════════
heading(doc, "10. Glossary")
gloss = [
    ("LangGraph",         "Open-source library for building stateful, graph-structured LLM workflows."),
    ("Node",              "A single processing step in the LangGraph graph (maps to one Python async function)."),
    ("Tool",              "A fine-grained async function that calls the LLM once and returns structured data."),
    ("CampaignReadinessState", "TypedDict shared across all nodes; the single source of truth for a workflow run."),
    ("Readiness Score",   "0–100 integer representing brief completeness; ≥ 85 bypasses the ambiguity step."),
    ("Session",           "An in-memory dict keyed by UUID representing one end-to-end workflow execution."),
    ("Mock Mode",         "Workflow execution without Azure OpenAI credentials; all tools return pre-built data."),
    ("Handoff Packet",    "Stakeholder-ready document containing executive summary, plan, risks, and next steps."),
    ("Uplift Simulation", "Statistical estimate of incremental campaign impact using historical segment data."),
    ("Assumption Register","List of AI-generated assumptions that fill brief gaps; requires manager approval."),
]
add_table(doc,
    ["Term", "Definition"],
    gloss,
    col_widths=[2.0, 5.75],
)

doc.add_paragraph()

# Footer note
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_footer.add_run(
    "Campaign Readiness Copilot  ·  Agent Interaction Reference  ·  Auto-generated"
)
run_f.font.size = Pt(8)
run_f.font.color.rgb = GRAY
run_f.italic = True

docx_path = os.path.join(OUT_DIR, "Interaction of agents.docx")
doc.save(docx_path)
print(f"Saved: {docx_path}")
